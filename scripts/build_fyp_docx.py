#!/usr/bin/env python3
"""
Build a print-ready FYP Word document with proper tables, spacing,
and clear INSERT-HERE zones for diagrams and screenshots.

Usage: python scripts/build_fyp_docx.py
Output: FYP-Final-Submission-IBBUL.docx (project root)
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
INPUT = ROOT / "FYP-Final-Submission-IBBUL.md"
OUTPUT = ROOT / "FYP-Final-Submission-IBBUL.docx"

FONT = "Times New Roman"
BODY_SIZE = Pt(12)
HEADING1_SIZE = Pt(14)
HEADING2_SIZE = Pt(12)
HEADING3_SIZE = Pt(12)
LINE_SPACING = 1.5
MARGIN = Inches(1)


def shade_cell(cell, fill: str = "FFF9C4") -> None:
    """Light yellow background for insert-here zones."""
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def set_cell_borders(cell) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "8")
        el.set(qn("w:color"), "666666")
        borders.append(el)
    tc_pr.append(borders)


def font_run(run, *, bold=False, italic=False, size=BODY_SIZE, color=None) -> None:
    run.font.name = FONT
    run.font.size = size
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    r = run._element
    r_pr = r.get_or_add_rPr()
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), FONT)
    r_fonts.set(qn("w:hAnsi"), FONT)
    r_fonts.set(qn("w:eastAsia"), FONT)
    r_pr.insert(0, r_fonts)


def format_paragraph(
    p,
    *,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    before=0,
    after=6,
    line=LINE_SPACING,
) -> None:
    pf = p.paragraph_format
    pf.alignment = align
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)


def add_styled_paragraph(
    doc,
    text: str = "",
    *,
    bold=False,
    italic=False,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    size=BODY_SIZE,
    before=0,
    after=6,
) -> None:
    p = doc.add_paragraph()
    format_paragraph(p, align=align, before=before, after=after)
    if text:
        run = p.add_run(text)
        font_run(run, bold=bold, italic=italic, size=size)


def add_rich_paragraph(doc, text: str, *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=6) -> None:
    """Paragraph with **bold** and *italic* spans."""
    p = doc.add_paragraph()
    format_paragraph(p, align=align, after=after)
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            font_run(run, bold=True)
        elif part.startswith("*") and part.endswith("*"):
            run = p.add_run(part[1:-1])
            font_run(run, italic=True)
        elif part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            font_run(run, italic=True)
        else:
            run = p.add_run(part)
            font_run(run)


def add_table(doc, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    table.autofit = True

    for ri, row in enumerate(rows):
        for ci in range(cols):
            cell = table.rows[ri].cells[ci]
            text = row[ci] if ci < len(row) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            format_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, after=3, before=3)
            run = p.add_run(text.strip())
            font_run(run, bold=(ri == 0))
            set_cell_borders(cell)

    doc.add_paragraph()  # spacing after table


def add_insert_zone(
    doc,
    figure_id: str,
    caption: str,
    instructions: str,
    source_hint: str = "",
) -> None:
    """Visible boxed area telling the student exactly what to insert."""
    doc.add_paragraph()
    banner = doc.add_paragraph()
    format_paragraph(banner, align=WD_ALIGN_PARAGRAPH.CENTER, before=12, after=6)
    run = banner.add_run(f"▼▼▼  INSERT {figure_id} BELOW THIS LINE  ▼▼▼")
    font_run(run, bold=True, size=Pt(11), color=RGBColor(0xC6, 0x28, 0x28))

    add_styled_paragraph(
        doc,
        caption,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=4,
    )
    if instructions:
        add_styled_paragraph(
            doc,
            instructions,
            italic=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            size=Pt(11),
            after=4,
        )
    if source_hint:
        add_styled_paragraph(
            doc,
            f"Source file: {source_hint}",
            align=WD_ALIGN_PARAGRAPH.CENTER,
            size=Pt(10),
            after=6,
        )

    # Image placeholder box
    box = doc.add_table(rows=1, cols=1)
    box.style = "Table Grid"
    cell = box.rows[0].cells[0]
    shade_cell(cell, "FFFDE7")
    set_cell_borders(cell)
    cell.width = Inches(5.5)

    lines = [
        "",
        "[ PASTE IMAGE OR SCREENSHOT HERE ]",
        "",
        f"Figure {figure_id}",
        "",
        "(Delete this grey text after inserting your image)",
        "",
        "",
        "",
    ]
    cell.text = "\n".join(lines)
    for p in cell.paragraphs:
        format_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, after=0, before=0)
        for r in p.runs:
            font_run(r, italic=True, size=Pt(11), color=RGBColor(0x75, 0x75, 0x75))

    cap = doc.add_paragraph()
    format_paragraph(cap, align=WD_ALIGN_PARAGRAPH.CENTER, before=6, after=12)
    run = cap.add_run(f"Figure {figure_id}: {caption}")
    font_run(run, bold=True, size=Pt(11))

    doc.add_paragraph()


FIGURE_SOURCES = {
    "3.1": ("System Architecture Diagram", "fyp-diagrams/01-system-architecture.mmd (export PNG at mermaid.live)"),
    "3.2": ("Entity-Relationship Diagram", "fyp-diagrams/02-er-diagram.mmd or FYP-Diagrams.drawio"),
    "3.3": ("Use Case Diagram", "fyp-diagrams/03-use-case.mmd"),
    "3.4": ("Activity Diagram — Event Approval", "fyp-diagrams/04-activity-approval.mmd"),
    "3.5": ("Data Flow Diagram — Event Creation", "fyp-diagrams/05-dfd-event-creation.mmd"),
    "4.1": ("Screenshot — Login Page", "fyp-diagrams/screenshots/Figure-4-1-Login.png  |  URL: /login"),
    "4.2": ("Screenshot — Dashboard", "fyp-diagrams/screenshots/Figure-4-2-Dashboard.png  |  Login: admin@nexus.dev"),
    "4.3": ("Screenshot — Event Creation", "fyp-diagrams/screenshots/Figure-4-3-Create-Event.png  |  Login: user@nexus.dev"),
    "4.4": ("Screenshot — Events Approval", "fyp-diagrams/screenshots/Figure-4-4-Events-Approve.png  |  Login: approver@nexus.dev"),
    "4.5": ("Screenshot — Calendar View", "fyp-diagrams/screenshots/Figure-4-5-Calendar.png  |  Login: approver@nexus.dev"),
    "4.6": ("Screenshot — Notifications Inbox", "fyp-diagrams/screenshots/Figure-4-6-Notifications.png  |  Login: approver@nexus.dev"),
    "4.7": ("Screenshot — Analytics Page", "fyp-diagrams/screenshots/Figure-4-7-Analytics.png  |  Login: admin@nexus.dev"),
    "4.8": ("Sequence Diagram — Event Approval (optional)", "fyp-diagrams/06-sequence-approve.mmd"),
}


def preprocess_markdown(raw: str) -> str:
    md = raw
    md = re.sub(
        r"^# FINAL YEAR PROJECT — PRINT & SUBMISSION PACKAGE[\s\S]*?^---\s*\n",
        "",
        md,
        count=1,
        flags=re.MULTILINE,
    )
    md = re.sub(r"<div align=\"center\">", "", md, flags=re.IGNORECASE)
    md = re.sub(r"</div>", "", md, flags=re.IGNORECASE)
    md = re.sub(r"<br\s*/?>", "\n\n", md, flags=re.IGNORECASE)
    md = md.replace("&nbsp;", " ")
    md = re.sub(r"\n{4,}", "\n\n\n", md)
    return md.strip()


def parse_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def is_table_sep(line: str) -> bool:
    return bool(re.match(r"^\|[\s\-:|]+\|\s*$", line.strip()))


def setup_document() -> Document:
    doc = Document()
    for section in doc.sections:
        section.top_margin = MARGIN
        section.bottom_margin = MARGIN
        section.left_margin = MARGIN
        section.right_margin = MARGIN

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = BODY_SIZE
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = LINE_SPACING
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    return doc


def build_document(md: str) -> Document:
    doc = setup_document()
    lines = md.splitlines()
    i = 0
    in_title = True  # until ## CERTIFICATION
    title_lines: list[str] = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip lone --- horizontal rules (add small space instead)
        if stripped == "---":
            if not in_title:
                doc.add_paragraph()
            i += 1
            continue

        # Code / mermaid blocks
        if stripped.startswith("```"):
            lang = stripped[3:].strip().lower()
            i += 1
            block: list[str] = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                block.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1  # closing ```

            # Figure mermaid — look back for figure caption
            if lang == "mermaid":
                fig_match = re.search(r"Figure\s+([\d.]+)", "\n".join(lines[max(0, i - 15) : i]))
                if fig_match:
                    fid = fig_match.group(1)
                    cap, src = FIGURE_SOURCES.get(fid, ("Diagram", "fyp-diagrams/"))
                    add_insert_zone(
                        doc,
                        fid,
                        cap,
                        "Export diagram as PNG (https://mermaid.live) and paste inside the box above.",
                        src,
                    )
            continue

        # Markdown table
        if stripped.startswith("|") and "|" in stripped[1:]:
            table_rows: list[list[str]] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                row_line = lines[i].strip()
                if not is_table_sep(row_line):
                    table_rows.append(parse_table_row(row_line))
                i += 1
            add_table(doc, table_rows)
            continue

        # Images → screenshot insert zone
        img = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", stripped)
        if img:
            alt, path = img.group(1), img.group(2)
            fig_m = re.search(r"Figure\s+([\d.]+)", alt) or re.search(r"Figure-([\d-]+)", path)
            if fig_m:
                fid = fig_m.group(1).replace("-", ".")
                if fid.startswith("4"):
                    fid = fid  # 4.1 etc
                cap, src = FIGURE_SOURCES.get(fid, (alt, path))
                add_insert_zone(
                    doc,
                    fid,
                    cap,
                    "Capture while app is running (pnpm dev). Password: ChangeMe123!",
                    src,
                )
            i += 1
            continue

        # Chapter headings — new page
        if re.match(r"^#\s+CHAPTER\s+", stripped, re.IGNORECASE):
            in_title = False
            doc.add_page_break()
            title = stripped.lstrip("#").strip()
            p = doc.add_paragraph()
            format_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=12)
            run = p.add_run(title.upper())
            font_run(run, bold=True, size=HEADING1_SIZE)
            i += 1
            continue

        # H1 (non-chapter)
        if stripped.startswith("# ") and not stripped.upper().startswith("# CHAPTER"):
            in_title = False
            text = stripped[2:].strip()
            if in_title or text.isupper():
                p = doc.add_paragraph()
                format_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, before=6, after=6)
                run = p.add_run(text)
                font_run(run, bold=True, size=HEADING1_SIZE)
            else:
                p = doc.add_paragraph()
                format_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, before=12, after=12)
                run = p.add_run(text)
                font_run(run, bold=True, size=HEADING1_SIZE)
            i += 1
            continue

        # H2
        if stripped.startswith("## "):
            in_title = False
            text = stripped[3:].strip()
            if text.upper() in ("CERTIFICATION", "DEDICATION", "ACKNOWLEDGEMENT", "ABSTRACT", "REFERENCES"):
                doc.add_page_break()
            p = doc.add_paragraph()
            format_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER if text.isupper() and len(text) < 40 else WD_ALIGN_PARAGRAPH.LEFT, before=12, after=6)
            run = p.add_run(text)
            font_run(run, bold=True, size=HEADING1_SIZE if text.isupper() and len(text) < 30 else HEADING2_SIZE)
            i += 1
            continue

        # H3
        if stripped.startswith("### "):
            in_title = False
            text = stripped[4:].strip()
            p = doc.add_paragraph()
            format_paragraph(p, before=10, after=4)
            run = p.add_run(text)
            font_run(run, bold=True, size=HEADING3_SIZE)
            i += 1
            continue

        # Figure caption line before mermaid (handled in block) — explicit **Figure X.X**
        fig_cap = re.match(r"^\*\*Figure\s+([\d.]+)\s+(.+?)\*\*\s*$", stripped)
        if fig_cap:
            fid, cap = fig_cap.group(1), fig_cap.group(2)
            # zone added when mermaid follows; if not, add now
            if i + 1 < len(lines) and lines[i + 1].strip().startswith("```"):
                i += 1
                continue
            src = FIGURE_SOURCES.get(fid, ("", ""))[1]
            add_insert_zone(doc, fid, cap, "Insert diagram image in the box above.", src)
            i += 1
            continue

        # Italic export note
        if stripped.startswith("*Export PNG") or stripped.startswith("*For Word"):
            add_styled_paragraph(doc, stripped.strip("*"), italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(10), after=4)
            i += 1
            continue

        # Bullet list
        if stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            format_paragraph(p, after=3)
            add_rich_inline(p, text)
            i += 1
            continue

        # Numbered list
        num = re.match(r"^(\d+)\.\s+(.+)", stripped)
        if num:
            p = doc.add_paragraph(style="List Number")
            format_paragraph(p, after=3)
            add_rich_inline(p, num.group(2))
            i += 1
            continue

        # Empty line
        if not stripped:
            if in_title:
                title_lines.append("")
            i += 1
            continue

        # Title page accumulation (before certification)
        if in_title:
            title_lines.append(stripped)
            i += 1
            continue

        # Signature lines
        if "Signature:" in stripped or stripped.startswith("Student "):
            add_styled_paragraph(doc, stripped, after=12)
            i += 1
            continue

        # Body paragraph
        add_rich_paragraph(doc, stripped)
        i += 1

    return doc


def add_rich_inline(p, text: str) -> None:
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            font_run(run, bold=True)
        elif part.startswith("*") and part.endswith("*"):
            run = p.add_run(part[1:-1])
            font_run(run, italic=True)
        elif part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            font_run(run, italic=True)
        else:
            run = p.add_run(part)
            font_run(run)


def build_title_page(doc: Document, title_lines: list[str]) -> None:
    """Re-build title page at start — called if we collected title lines."""
    pass  # title handled inline in main loop via in_title; simplify by preprocessing


def main() -> int:
    if not INPUT.exists():
        print(f"Missing input: {INPUT}", file=sys.stderr)
        return 1

    raw = INPUT.read_text(encoding="utf-8")
    md = preprocess_markdown(raw)

    # Dedicated title page before main content
    doc = setup_document()

    # --- Title page ---
    title_block = [
        "",
        "DESIGN AND IMPLEMENTATION OF A UNIVERSITY EVENT SCHEDULING AND NOTIFICATION SYSTEM",
        "",
        "A Case Study of Ibrahim Badamasi Babangida University, Lapai",
        "",
        "",
        "BY",
        "",
        "[STUDENT ONE FULL NAME]",
        "Matriculation Number: [UG__/CS/____]",
        "",
        "AND",
        "",
        "[STUDENT TWO FULL NAME]",
        "Matriculation Number: [UG__/CS/____]",
        "",
        "",
        "A Final Year Project Submitted to the Department of Computer Science,",
        "Faculty of Science,",
        "Ibrahim Badamasi Babangida University, Lapai,",
        "Niger State, Nigeria.",
        "",
        "",
        "In Partial Fulfilment of the Requirements for the Award of the",
        "Bachelor of Science (B.Sc.) Degree in Computer Science",
        "",
        "",
        "[MONTH, YEAR]",
    ]
    for t in title_block:
        p = doc.add_paragraph()
        format_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=8)
        if t:
            run = p.add_run(t)
            bold = t.isupper() and len(t) > 20 and "BY" not in t and "AND" not in t
            font_run(run, bold=bold, size=HEADING1_SIZE if bold else BODY_SIZE)

    doc.add_page_break()

    # Parse body (skip title section in md — from CERTIFICATION onward)
    cert_idx = md.find("## CERTIFICATION")
    body = md[cert_idx:] if cert_idx >= 0 else md

    # Re-use build logic on body only
    lines = body.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped == "---":
            doc.add_paragraph()
            i += 1
            continue

        if stripped.startswith("```"):
            lang = stripped[3:].strip().lower()
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                i += 1
            if i < len(lines):
                i += 1
            if lang == "mermaid":
                ctx = "\n".join(lines[max(0, i - 12) : min(len(lines), i)])
                fig_match = re.search(r"Figure\s+([\d.]+)[:\s]+([^\n*]+)", ctx)
                if fig_match:
                    fid = fig_match.group(1)
                    cap = fig_match.group(2).strip()
                    src = FIGURE_SOURCES.get(fid, (cap, "fyp-diagrams/"))[1]
                    add_insert_zone(
                        doc,
                        fid,
                        cap,
                        "Export as PNG from https://mermaid.live and paste in the yellow box.",
                        src,
                    )
            continue

        if stripped.startswith("|") and "|" in stripped[1:]:
            table_rows: list[list[str]] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                row_line = lines[i].strip()
                if not is_table_sep(row_line):
                    table_rows.append(parse_table_row(row_line))
                i += 1
            add_table(doc, table_rows)
            continue

        img = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", stripped)
        if img:
            alt, path = img.group(1), img.group(2)
            fig_m = re.search(r"Figure\s+([\d.]+)", alt)
            fid = fig_m.group(1) if fig_m else "?"
            cap, src = FIGURE_SOURCES.get(fid, (alt, path))
            add_insert_zone(
                doc,
                fid,
                cap,
                "Run: pnpm dev → open browser → capture screenshot → paste in yellow box.",
                src,
            )
            i += 1
            continue

        if re.match(r"^#\s+CHAPTER\s+", stripped, re.IGNORECASE):
            doc.add_page_break()
            title = stripped.lstrip("#").strip()
            p = doc.add_paragraph()
            format_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=12)
            run = p.add_run(title.upper())
            font_run(run, bold=True, size=HEADING1_SIZE)
            i += 1
            continue

        if stripped.startswith("## "):
            text = stripped[3:].strip()
            if text.upper() in ("CERTIFICATION", "DEDICATION", "ACKNOWLEDGEMENT", "ABSTRACT", "REFERENCES", "APPENDICES"):
                if text.upper() != "CERTIFICATION":
                    doc.add_page_break()
            p = doc.add_paragraph()
            center = text.upper() in (
                "CERTIFICATION",
                "DEDICATION",
                "ACKNOWLEDGEMENT",
                "ABSTRACT",
                "TABLE OF CONTENTS",
                "LIST OF TABLES",
                "LIST OF FIGURES",
                "REFERENCES",
                "APPENDICES",
            )
            format_paragraph(
                p,
                align=WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT,
                before=12,
                after=6,
            )
            run = p.add_run(text)
            font_run(run, bold=True, size=HEADING1_SIZE if center else HEADING2_SIZE)
            i += 1
            continue

        if stripped.startswith("### "):
            text = stripped[4:].strip()
            p = doc.add_paragraph()
            format_paragraph(p, before=10, after=4)
            run = p.add_run(text)
            font_run(run, bold=True, size=HEADING3_SIZE)
            i += 1
            continue

        fig_cap = re.match(r"^\*\*Figure\s+([\d.]+)\s+(.+?)\*\*\s*$", stripped)
        if fig_cap:
            if i + 1 < len(lines) and lines[i + 1].strip().startswith("```"):
                i += 1
                continue
            fid, cap = fig_cap.group(1), fig_cap.group(2)
            src = FIGURE_SOURCES.get(fid, ("", ""))[1]
            add_insert_zone(doc, fid, cap, "Insert image in the yellow box above.", src)
            i += 1
            continue

        if stripped.startswith("*Export") or stripped.startswith("*For Word"):
            i += 1
            continue

        if stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            format_paragraph(p, after=3)
            add_rich_inline(p, stripped[2:].strip())
            i += 1
            continue

        num = re.match(r"^(\d+)\.\s+(.+)", stripped)
        if num:
            p = doc.add_paragraph(style="List Number")
            format_paragraph(p, after=3)
            add_rich_inline(p, num.group(2))
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if "Keywords:" in stripped:
            add_rich_paragraph(doc, stripped, after=12)
            i += 1
            continue

        add_rich_paragraph(doc, stripped)
        i += 1

    doc.save(OUTPUT)
    size_kb = OUTPUT.stat().st_size / 1024
    print(f"Created: {OUTPUT}")
    print(f"Size: {size_kb:.1f} KB")
    print("Tables: bordered grid | Figures: yellow INSERT boxes | Font: Times New Roman 12pt")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
